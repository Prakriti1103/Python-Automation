from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import difflib, re, os

# ========================== Helper Functions ==========================

def get_alignment(para):
    return {
        WD_ALIGN_PARAGRAPH.LEFT: "LEFT",
        WD_ALIGN_PARAGRAPH.CENTER: "CENTER",
        WD_ALIGN_PARAGRAPH.RIGHT: "RIGHT",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "JUSTIFY",
        None: "UNKNOWN"
    }.get(para.alignment, "UNKNOWN")

def get_indent(para):
    l = para.paragraph_format.left_indent
    r = para.paragraph_format.right_indent
    return (round(l.cm, 2) if l else 0.0, round(r.cm, 2) if r else 0.0)

def get_line_spacing(para):
    return round(para.paragraph_format.line_spacing, 2) if para.paragraph_format.line_spacing else "Default"

def extract_paragraph_data(doc):
    data = []
    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        run = para.runs[0] if para.runs else None
        data.append({
            "text": para.text.strip(),
            "font_name": run.font.name if run and run.font.name else "Default",
            "font_size": run.font.size.pt if run and run.font.size else "Default",
            "bold": run.bold if run else False,
            "italic": run.italic if run else False,
            "alignment": get_alignment(para),
            "spacing": get_line_spacing(para),
            "left_indent": get_indent(para)[0],
            "right_indent": get_indent(para)[1]
        })
    return data

def get_cell_style(cell):
    style = {
        "text": cell.text.strip(),
        "font_name": "Default",
        "font_size": "Default",
        "bold": False,
        "italic": False,
        "alignment": "UNKNOWN"
    }
    if cell.paragraphs:
        para = cell.paragraphs[0]
        run = para.runs[0] if para.runs else None
        if run:
            style.update({
                "font_name": run.font.name or "Default",
                "font_size": run.font.size.pt if run.font.size else "Default",
                "bold": run.bold or False,
                "italic": run.italic or False,
                "alignment": get_alignment(para)
            })
    return style

def get_word_diff(old, new):
    matcher = difflib.SequenceMatcher(None, old.split(), new.split())
    diffs = []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'replace':
            diffs += [f"[Word Changed] {old.split()[i]} → {new.split()[j]}" for i, j in zip(range(i1, i2), range(j1, j2))]
        elif tag == 'insert':
            diffs += [f"[Inserted] {new.split()[j]}" for j in range(j1, j2)]
        elif tag == 'delete':
            diffs += [f"[Deleted] {old.split()[i]}" for i in range(i1, i2)]
    return diffs

def detect_spacing_issues(old, new):
    issues = []
    if re.search(r"\s{2,}", old) and not re.search(r"\s{2,}", new):
        issues.append(f"[Extra Spaces Detected] \"{old}\" → \"{new}\"")
    if re.sub(r'\s+', '', old) == re.sub(r'\s+', '', new) and old != new:
        issues.append(f"[Letter Spacing Issue] \"{old}\" → \"{new}\"")
    return issues

def add_colored_paragraph(doc, label, content, color):
    p = doc.add_paragraph()
    run = p.add_run(f"{label} {content}")
    run.font.color.rgb = RGBColor(*color)

# ========================== Comparison Sections ==========================

def compare_paragraphs(pre, post, doc_out):
    doc_out.add_heading("Paragraph Comparison", level=1)
    for i in range(max(len(pre), len(post))):
        preview = pre[i]['text'][:100] + "..." if i < len(pre) else post[i]['text'][:100] + "..."
        doc_out.add_paragraph(f"Line {i+1}: {preview}", style="Heading3")
        if i >= len(pre):
            add_colored_paragraph(doc_out, "[Added]", post[i]['text'], (0, 128, 0)); continue
        if i >= len(post):
            add_colored_paragraph(doc_out, "[Removed]", pre[i]['text'], (255, 0, 0)); continue
        a, b = pre[i], post[i]
        if a['text'] != b['text']:
            add_colored_paragraph(doc_out, "[Text Changed] Old:", a['text'], (255, 0, 0))
            add_colored_paragraph(doc_out, "New:", b['text'], (0, 100, 0))
            for diff in get_word_diff(a['text'], b['text']) + detect_spacing_issues(a['text'], b['text']):
                add_colored_paragraph(doc_out, "", diff, (255, 165, 0))
        for k in ['font_name','font_size','bold','italic','alignment','spacing','left_indent','right_indent']:
            if a[k] != b[k]:
                add_colored_paragraph(doc_out, "", f"[{k.replace('_',' ').title()} Changed] {a[k]} → {b[k]}", (255, 165, 0))

def compare_table_styles(doc1, doc2, doc_out):
    doc_out.add_heading("Table Comparison", level=1)
    t1, t2 = doc1.tables, doc2.tables
    for t in range(max(len(t1), len(t2))):
        if t >= len(t1):
            doc_out.add_paragraph(f"Table {t+1}: (Added)", style="Heading3")
            continue
        if t >= len(t2):
            doc_out.add_paragraph(f"Table {t+1}: (Removed)", style="Heading3")
            continue
        header_cells = t1[t].rows[0].cells if t1[t].rows else []
        headers = " | ".join([cell.text.strip() for cell in header_cells])
        doc_out.add_paragraph(f"Table {t+1}: ({headers})", style="Heading3")
        rows1, rows2 = t1[t].rows, t2[t].rows
        for r in range(max(len(rows1), len(rows2))):
            if r >= len(rows1) or r >= len(rows2): continue
            cells1, cells2 = rows1[r].cells, rows2[r].cells
            for c in range(max(len(cells1), len(cells2))):
                if c >= len(cells1) or c >= len(cells2): continue
                c1, c2 = get_cell_style(cells1[c]), get_cell_style(cells2[c])
                if c1["text"] != c2["text"]:
                    add_colored_paragraph(doc_out, f"R{r+1}C{c+1} Old:", c1["text"], (255, 0, 0))
                    add_colored_paragraph(doc_out, "New:", c2["text"], (0, 100, 0))
                    for diff in get_word_diff(c1["text"], c2["text"]) + detect_spacing_issues(c1["text"], c2["text"]):
                        add_colored_paragraph(doc_out, "", diff, (255, 165, 0))
                for attr in ['font_name','font_size','bold','italic','alignment']:
                    if c1[attr] != c2[attr]:
                        add_colored_paragraph(doc_out, "", f"[{attr.replace('_',' ').title()} Changed] {c1[attr]} → {c2[attr]}", (255, 165, 0))

def compare_images(doc1, doc2, doc_out):
    doc_out.add_heading("Image Comparison", level=1)
    shapes1 = doc1.inline_shapes
    shapes2 = doc2.inline_shapes
    max_len = max(len(shapes1), len(shapes2))
    for i in range(max_len):
        if i >= len(shapes1):
            add_colored_paragraph(doc_out, f"[Image Added]", f"Image {i+1} (Width: {shapes2[i].width} EMU, Height: {shapes2[i].height} EMU)", (0, 128, 0))
            continue
        if i >= len(shapes2):
            add_colored_paragraph(doc_out, f"[Image Removed]", f"Image {i+1} (Width: {shapes1[i].width} EMU, Height: {shapes1[i].height} EMU)", (255, 0, 0))
            continue
        if shapes1[i].width != shapes2[i].width or shapes1[i].height != shapes2[i].height:
            add_colored_paragraph(doc_out, f"Image {i+1}:", "", (0, 0, 255))
            if shapes1[i].width != shapes2[i].width:
                add_colored_paragraph(doc_out, "", f"[Width Changed] {shapes1[i].width} → {shapes2[i].width} EMU", (255, 165, 0))
            if shapes1[i].height != shapes2[i].height:
                add_colored_paragraph(doc_out, "", f"[Height Changed] {shapes1[i].height} → {shapes2[i].height} EMU", (255, 165, 0))

# ========================== Main Execution ==========================

def main():
    doc1 = Document("May 19 buck.docx")
    doc2 = Document("May 28 buck.docx")
    output_doc = Document()

    compare_paragraphs(extract_paragraph_data(doc1), extract_paragraph_data(doc2), output_doc)
    compare_table_styles(doc1, doc2, output_doc)
    compare_images(doc1, doc2, output_doc)

    output_doc.save("final_doc_comparison_with_images.docx")
    print(" Done! Output saved as: final_doc_comparison_with_images.docx")

if __name__ == "__main__":
    main()
