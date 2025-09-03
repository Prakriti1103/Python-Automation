import streamlit as st
from docx import Document
from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import difflib, re
from zipfile import ZipFile
from lxml import etree
from io import BytesIO
from difflib import SequenceMatcher, get_close_matches
import zipfile
import xml.etree.ElementTree as ET
import math
from docx.text.paragraph import Paragraph
from docx.oxml import parse_xml
import hashlib
from docx.enum.text import WD_BREAK

# === UTILS ===
# Add a paragraph to the output document with a colored label and content
def add_colored_paragraph(doc, label, content, color):
    p = doc.add_paragraph()
    run = p.add_run(f"{label} {content}")
    run.font.color.rgb = RGBColor(*color)

def get_alignment(para):
    return {
        WD_ALIGN_PARAGRAPH.LEFT: "LEFT",
        WD_ALIGN_PARAGRAPH.CENTER: "CENTER",
        WD_ALIGN_PARAGRAPH.RIGHT: "RIGHT",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "JUSTIFY",
        None: "UNKNOWN"
    }.get(para.alignment, "UNKNOWN")

def get_indent(para):
    l = para.paragraph_format.left_indent
    r = para.paragraph_format.right_indent
    return (round(l.cm, 2) if l else 0.0, round(r.cm, 2) if r else 0.0)

def get_line_spacing(para):
    return round(para.paragraph_format.line_spacing, 2) if para.paragraph_format.line_spacing else "Default"

def get_highlight_name(color_enum):
    highlight_map = {
        WD_COLOR_INDEX.AUTO: "auto",
        WD_COLOR_INDEX.BLACK: "black",
        WD_COLOR_INDEX.BLUE: "blue",
        WD_COLOR_INDEX.BRIGHT_GREEN: "bright green",
        WD_COLOR_INDEX.DARK_BLUE: "dark blue",
        WD_COLOR_INDEX.DARK_RED: "dark red",
        WD_COLOR_INDEX.DARK_YELLOW: "dark yellow",
        WD_COLOR_INDEX.GRAY_25: "gray 25%",
        WD_COLOR_INDEX.GRAY_50: "gray 50%",
        WD_COLOR_INDEX.GREEN: "green",
        WD_COLOR_INDEX.PINK: "pink",
        WD_COLOR_INDEX.RED: "red",
        WD_COLOR_INDEX.TEAL: "teal",
        WD_COLOR_INDEX.TURQUOISE: "turquoise",
        WD_COLOR_INDEX.VIOLET: "violet",
        WD_COLOR_INDEX.WHITE: "white",
        WD_COLOR_INDEX.YELLOW: "yellow",
        None: "None"
    }
    return highlight_map.get(color_enum, str(color_enum))

def get_rgb_color_name(rgb):
    rgb_map = {
        RGBColor(255, 0, 0): "red",
        RGBColor(0, 255, 0): "green",
        RGBColor(0, 0, 255): "blue",
        RGBColor(255, 255, 0): "yellow",
        RGBColor(0, 0, 0): "black",
        RGBColor(255, 255, 255): "white",
        RGBColor(128, 128, 128): "gray",
        RGBColor(0, 255, 255): "cyan",
        RGBColor(255, 0, 255): "magenta",
        None: "None"
    }
    return rgb_map.get(rgb, str(rgb))


PARA_SIMILARITY_FOR_RUNLEVEL = 0.60  # for 'replace' paragraphs, only check run-level if overall para text is somewhat similar

def get_run_style(run):
    """Return normalized run-level style values (strings and booleans)."""
    if run is None:
        return {
            'text': '',
            'bold': False, 'italic': False, 'underline': False,
            'font_name': "Default", 'font_size': "Default",
            'font_color': "None", 'highlight': "None"
        }

    # convert font color to friendly name (use existing mapping function)
    try:
        raw_color = run.font.color.rgb if run.font.color and run.font.color.rgb else None
    except Exception:
        raw_color = None
    try:
        raw_highlight = run.font.highlight_color if run.font.highlight_color else None
    except Exception:
        raw_highlight = None

    return {
        'text': run.text or '',
        'bold': bool(run.bold),
        'italic': bool(run.italic),
        'underline': bool(run.underline),
        'font_name': run.font.name if run.font.name else "Default",
        'font_size': run.font.size.pt if run.font.size else "Default",
        'font_color': get_rgb_color_name(raw_color),
        'highlight': get_highlight_name(raw_highlight)
    }

def get_paragraph_info(para):
    run = para.runs[0] if para.runs else None
    
    # Helper maps
    def get_highlight_name(color_enum):
        highlight_map = {
            WD_COLOR_INDEX.AUTO: "auto",
            WD_COLOR_INDEX.BLACK: "black",
            WD_COLOR_INDEX.BLUE: "blue",
            WD_COLOR_INDEX.BRIGHT_GREEN: "bright green",
            WD_COLOR_INDEX.DARK_BLUE: "dark blue",
            WD_COLOR_INDEX.DARK_RED: "dark red",
            WD_COLOR_INDEX.DARK_YELLOW: "dark yellow",
            WD_COLOR_INDEX.GRAY_25: "gray 25%",
            WD_COLOR_INDEX.GRAY_50: "gray 50%",
            WD_COLOR_INDEX.GREEN: "green",
            WD_COLOR_INDEX.PINK: "pink",
            WD_COLOR_INDEX.RED: "red",
            WD_COLOR_INDEX.TEAL: "teal",
            WD_COLOR_INDEX.TURQUOISE: "turquoise",
            WD_COLOR_INDEX.VIOLET: "violet",
            WD_COLOR_INDEX.WHITE: "white",
            WD_COLOR_INDEX.YELLOW: "yellow",
            None: "None"
        }
        return highlight_map.get(color_enum, str(color_enum))

    def get_rgb_color_name(rgb):
        rgb_map = {
            RGBColor(255, 0, 0): "red",
            RGBColor(0, 255, 0): "green",
            RGBColor(0, 0, 255): "blue",
            RGBColor(255, 255, 0): "yellow",
            RGBColor(0, 0, 0): "black",
            RGBColor(255, 255, 255): "white",
            RGBColor(128, 128, 128): "gray",
            RGBColor(0, 255, 255): "cyan",
            RGBColor(255, 0, 255): "magenta",
            None: "None"
        }
        return rgb_map.get(rgb, str(rgb))

    return {
        "text": para.text.strip(),
        "font_name": run.font.name if run and run.font.name else "Default",
        "font_size": run.font.size.pt if run and run.font.size else "Default",
        "bold": run.bold if run else False,
        "italic": run.italic if run else False,
        "underline": run.underline if run else False,
        "highlight": get_highlight_name(run.font.highlight_color) if run else "None",
        "font_color": get_rgb_color_name(run.font.color.rgb if run and run.font.color and run.font.color.rgb else None),
        "alignment": get_alignment(para),
        "spacing": get_line_spacing(para),
        "left_indent": get_indent(para)[0],
        "right_indent": get_indent(para)[1]
    }

def compare_run_styles(r1, r2):
    """Compare two run-style dicts and return differences dict."""
    keys = ['bold', 'italic', 'underline', 'font_name', 'font_size', 'font_color', 'highlight', 'text']
    diffs = {}
    for k in keys:
        v1 = r1.get(k, "Default")
        v2 = r2.get(k, "Default")
        # normalize None to Default for font attributes
        if k in ('font_name', 'font_size', 'font_color', 'highlight'):
            if v1 is None: v1 = "Default"
            if v2 is None: v2 = "Default"
        if v1 != v2:
            diffs[k] = (v1, v2)
    return diffs

def _compare_format_dicts(d1, d2):
    
    keys = ['font_name','font_size','bold','italic','underline','font_color','highlight','alignment','spacing','left_indent','right_indent']
    diffs = {}
    for k in keys:
        v1 = d1.get(k, "Default")
        v2 = d2.get(k, "Default")
        # Normalize some common equivalents
        if v1 is None: v1 = "Default"
        if v2 is None: v2 = "Default"
        if v1 != v2:
            diffs[k] = (v1, v2)
    return diffs


def get_word_diff(old, new):
    old_words, new_words = old.split(), new.split()
    matcher = difflib.SequenceMatcher(None, old_words, new_words)
    diffs = []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag in ['replace', 'insert', 'delete']:
            ow = " ".join(old_words[i1:i2])
            nw = " ".join(new_words[j1:j2])
            if ow != nw:
                diffs.append(f"[Word Changed] {ow} → {nw}")
    return diffs

def get_letter_spacing_from_run(run):
    try:
        xml = run._element
        spacing_elem = xml.find('.//w:spacing', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
        if spacing_elem is not None and 'w:val' in spacing_elem.attrib:
            return spacing_elem.attrib['w:val']
    except:
        pass
    return "Default"

def detect_spacing_issues(old, new):
    issues = []
    if re.search(r"\s{2,}", old) and not re.search(r"\s{2,}", new):
        issues.append(f"[Extra Spaces Detected] \"{old}\" → \"{new}\"")
    if re.sub(r'\s+', '', old) == re.sub(r'\s+', '', new) and old != new:
        issues.append(f"[Letter Spacing Issue] \"{old}\" → \"{new}\"")
    return issues

def fuzzy_compare_lists(list1, list2, label, output_doc):
    matcher = difflib.SequenceMatcher(None, list1, list2)
    table_count = 0
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            continue
        elif tag == 'replace':
            for i, j in zip(range(i1, i2), range(j1, j2)):
                preview = list2[j][:80] + ("..." if len(list2[j]) > 80 else "")
                heading_label = f"{label} {j+1}: \"{preview}\""
                if label == "TableRow":
                    table_count += 1
                    heading_label = f"Table {table_count} | Preview: \"{preview}\""
                output_doc.add_paragraph(heading_label, style="Heading 3")
                add_colored_paragraph(output_doc, f"[Old {label} {i+1}]", list1[i], (255, 0, 0))
                add_colored_paragraph(output_doc, f"[New {label} {j+1}]", list2[j], (0, 128, 0))
                for diff in get_word_diff(list1[i], list2[j]) + detect_spacing_issues(list1[i], list2[j]):
                    add_colored_paragraph(output_doc, "", diff, (255, 165, 0))
        elif tag == 'delete':
            for i in range(i1, i2):
                output_doc.add_paragraph(f"Removed {label} (Line {i+1})", style="Heading 3")
                add_colored_paragraph(output_doc, f"[Removed {label}]", list1[i], (255, 0, 0))
        elif tag == 'insert':
            for j in range(j1, j2):
                preview = list2[j][:80] + ("..." if len(list2[j]) > 80 else "")
                heading_label = f"Added {label} (Line {j+1})"
                if label == "TableRow":
                    table_count += 1
                    heading_label = f"Table {table_count} | Preview: \"{preview}\""
                output_doc.add_paragraph(heading_label, style="Heading 3")
                add_colored_paragraph(output_doc, f"[Added {label}]", list2[j], (0, 128, 0))

def estimate_paragraph_pages(doc):
    avg_paragraphs_per_page = 25
    para_map = []
    for i, para in enumerate(doc.paragraphs):
        page = (i // avg_paragraphs_per_page) + 1
        para_map.append(page)
    return para_map

# ===================================== Paragraph Comparison =================================

# ---- tiny helpers (scoped inside; no external deps) ----
rgb_name_map = {
        (255, 0, 0): "red", (0, 255, 0): "green", (0, 0, 255): "blue",
        (255, 255, 0): "yellow", (0, 0, 0): "black", (255, 255, 255): "white",
        (128, 128, 128): "gray", (0, 255, 255): "cyan", (255, 0, 255): "magenta"
    }
    
def color_to_name(rgb):
        if rgb is None:
            return "None"
        try:
            tup = (int(rgb[0]), int(rgb[1]), int(rgb[2]))
            return rgb_name_map.get(tup, f"rgb{tup}")
        except Exception:
            return str(rgb) if rgb is not None else "None"

def highlight_to_name(h):
        table = {
            WD_COLOR_INDEX.AUTO: "auto", WD_COLOR_INDEX.BLACK: "black",
            WD_COLOR_INDEX.BLUE: "blue", WD_COLOR_INDEX.BRIGHT_GREEN: "bright green",
            WD_COLOR_INDEX.DARK_BLUE: "dark blue", WD_COLOR_INDEX.DARK_RED: "dark red",
            WD_COLOR_INDEX.DARK_YELLOW: "dark yellow", WD_COLOR_INDEX.GRAY_25: "gray 25%",
            WD_COLOR_INDEX.GRAY_50: "gray 50%", WD_COLOR_INDEX.GREEN: "green",
            WD_COLOR_INDEX.PINK: "pink", WD_COLOR_INDEX.RED: "red",
            WD_COLOR_INDEX.TEAL: "teal", WD_COLOR_INDEX.TURQUOISE: "turquoise",
            WD_COLOR_INDEX.VIOLET: "violet", WD_COLOR_INDEX.WHITE: "white",
            WD_COLOR_INDEX.YELLOW: "yellow", None: "None"
        }
        return table.get(h, str(h))

def get_alignment_name(para):
        return {
            WD_ALIGN_PARAGRAPH.LEFT: "LEFT",
            WD_ALIGN_PARAGRAPH.CENTER: "CENTER",
            WD_ALIGN_PARAGRAPH.RIGHT: "RIGHT",
            WD_ALIGN_PARAGRAPH.JUSTIFY: "JUSTIFY",
            None: "UNKNOWN"
        }.get(para.alignment, "UNKNOWN")

def get_spacing(para):
        ls = para.paragraph_format.line_spacing
        try:
            return round(float(ls), 2) if ls else "Default"
        except Exception:
            return "Default"

def get_indent_cm(para):
        pf = para.paragraph_format
        l = pf.left_indent.cm if pf.left_indent else 0.0
        r = pf.right_indent.cm if pf.right_indent else 0.0
        return round(l, 2), round(r, 2)

def dominant_run_style(para):
        # weight by number of characters per run
        counts = {
            "bold": {}, "italic": {}, "underline": {},
            "font_name": {}, "font_size": {},
            "font_color": {}, "highlight": {}
        }
        for run in para.runs:
            txt = run.text or ""
            n = len(txt)
            if n == 0:
                continue

            def bump(k, v):
                counts[k][v] = counts[k].get(v, 0) + n

            b = True if run.bold else False
            i = True if run.italic else False
            u = True if bool(run.underline) else False
            fname = run.font.name if run.font and run.font.name else "Default"
            fsize = float(run.font.size.pt) if run.font and run.font.size else "Default"
            fcol = color_to_name(run.font.color.rgb if run.font and run.font.color and run.font.color.rgb else None)
            hi = highlight_to_name(run.font.highlight_color if run.font and run.font.highlight_color else None)

            bump("bold", b); bump("italic", i); bump("underline", u)
            bump("font_name", fname); bump("font_size", fsize)
            bump("font_color", fcol); bump("highlight", hi)

        def pick(k, default):
            bucket = counts[k]
            return max(bucket.items(), key=lambda kv: kv[1])[0] if bucket else default

        style = {
            "bold": pick("bold", False),
            "italic": pick("italic", False),
            "underline": pick("underline", False),
            "font_name": pick("font_name", "Default"),
            "font_size": pick("font_size", "Default"),
            "font_color": pick("font_color", "None"),
            "highlight": pick("highlight", "None"),
            "alignment": get_alignment_name(para),
            "spacing": get_spacing(para)
        }
        l, r = get_indent_cm(para)
        style["left_indent"] = l
        style["right_indent"] = r
        return style

def safe_changed(key, a, b):
        # Reduce false-positives: only flag when both sides are explicit or a real boolean flip
        if key in ("font_name", "font_color", "highlight"):
            if a in ("Default", "None") and b in ("Default", "None"):
                return False
        if key in ("font_size", "left_indent", "right_indent", "spacing"):
            if a == "Default" or b == "Default":
                return False
        return a != b

def style_diffs(pre, post):
        keys = [
            "alignment", "spacing", "left_indent", "right_indent",
            "bold", "italic", "underline",
            "font_name", "font_size", "font_color", "highlight"
        ]
        diffs = []
        for k in keys:
            if safe_changed(k, pre[k], post[k]):
                diffs.append((k, pre[k], post[k]))
        return diffs
    # ---- end helpers ----
    
def compare_paragraphs(doc1, doc2, output_doc):
    from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH
    import difflib

    # Keep paragraph objects aligned with their stripped text
    paras1_objs = [p for p in doc1.paragraphs if p.text.strip()]
    paras2_objs = [p for p in doc2.paragraphs if p.text.strip()]
    paras1 = [p.text.strip() for p in paras1_objs]
    paras2 = [p.text.strip() for p in paras2_objs]

    output_doc.add_heading("Paragraph Comparison", level=1)
    sm = difflib.SequenceMatcher(None, paras1, paras2)

    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            # Text is identical — show formatting-only changes (using dominant style)
            for idx in range(0, min(i2 - i1, j2 - j1)):
                if (i1 + idx) >= len(paras1_objs) or (j1 + idx) >= len(paras2_objs):
                    continue
                p1 = paras1_objs[i1 + idx]
                p2 = paras2_objs[j1 + idx]
                d1 = dominant_run_style(p1)
                d2 = dominant_run_style(p2)
                diffs = style_diffs(d1, d2)
                if diffs:
                    preview = paras2[j1 + idx][:80] + ("..." if len(paras2[j1 + idx]) > 80 else "")
                    output_doc.add_paragraph(f"Paragraph {j1 + idx + 1} | Preview: \"{preview}\"", style="Heading 3")
                    for k, a, b in diffs:
                        label = f"[{k.replace('_', ' ').title()} Changed]"
                        add_colored_paragraph(output_doc, label, f"{a} → {b}", (0, 0, 255))
            continue

        # For replace/delete/insert, keep your original behavior + add safe style diffs on replace
        if tag == "replace":
            pairs = list(zip(range(i1, i2), range(j1, j2)))
        elif tag == "delete":
            pairs = [(i, None) for i in range(i1, i2)]
        else:  # insert
            pairs = [(None, j) for j in range(j1, j2)]

        for i, j in pairs:
            old = paras1[i] if i is not None else None
            new = paras2[j] if j is not None else None

            if i is not None and j is not None:
                preview = new[:80] + ("..." if len(new) > 80 else "")
                output_doc.add_paragraph(f"Paragraph {j + 1} | Preview: \"{preview}\"", style="Heading 3")
                add_colored_paragraph(output_doc, "[Old Paragraph]", old, (255, 0, 0))
                add_colored_paragraph(output_doc, "[New Paragraph]", new, (0, 128, 0))

                # word changes + spacing issues (your existing helpers)
                for diff in get_word_diff(old, new) + detect_spacing_issues(old, new):
                    add_colored_paragraph(output_doc, "", diff, (255, 165, 0))

                # dominant style diffs (safe)
                if i < len(paras1_objs) and j < len(paras2_objs):
                    d1 = dominant_run_style(paras1_objs[i])
                    d2 = dominant_run_style(paras2_objs[j])
                    for k, a, b in style_diffs(d1, d2):
                        label = f"[{k.replace('_', ' ').title()} Changed]"
                        add_colored_paragraph(output_doc, label, f"{a} → {b}", (0, 0, 255))

            elif i is not None:
                output_doc.add_paragraph(f"Removed Paragraph (Line {i + 1})", style="Heading 3")
                add_colored_paragraph(output_doc, "[Removed Paragraph]", old, (255, 0, 0))

            elif j is not None:
                preview = new[:80] + ("..." if len(new) > 80 else "")
                output_doc.add_paragraph(f"Added Paragraph {j + 1} | Preview: \"{preview}\"", style="Heading 3")
                add_colored_paragraph(output_doc, "[Added Paragraph]", new, (0, 128, 0))

# ======================================= Textbox Comparison ============================================
def extract_textbox_paragraphs_with_pages(docx_path):
    paragraphs_with_pages = []
    try:
        with ZipFile(docx_path) as docx:
            xml_content = docx.read("word/document.xml")
            tree = etree.fromstring(xml_content)
            namespaces = {
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape'
            }
            page_number = 1
            count = 0
            for para in tree.xpath('.//w:drawing//wps:txbx//w:p', namespaces=namespaces):
                texts = para.xpath('.//w:t', namespaces=namespaces)
                text_content = "".join([t.text for t in texts if t.text])
                if text_content.strip():
                    count += 1
                    page_number = (count // 5) + 1  # assume 5 textboxes per page avg
                    paragraphs_with_pages.append((text_content.strip(), page_number))
    except:
        pass
    return paragraphs_with_pages

def get_textbox_estimated_page(doc, idx):
    
    # Very naive: index // 5 ≈ page grouping, adjust if needed
    return (idx // 5) + 1


# --- constants ---
EMU_PER_INCH = 914400  # Word stores shape sizes in EMUs


# --- helpers for units ---
def emu_to_inches(emu_val):
    try:
        return round(int(emu_val) / EMU_PER_INCH, 2)
    except Exception:
        return None


# --- extractor ---
def extract_textboxes(doc):
   
    EMU_PER_INCH = 914400
    def emu_to_inches(emu_val):
        try:
            return round(int(emu_val) / EMU_PER_INCH, 2)
        except Exception:
            return None

    TXBX = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}txbxContent"
    PARA = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p"
    EXT  = "{http://schemas.openxmlformats.org/drawingml/2006/main}ext"

    seen = set()
    textboxes = []

    for shape in doc.element.iter(TXBX):
        paras = []
        for p in shape.iter(PARA):
            try:
                paras.append(Paragraph(p, doc))
            except Exception:
                pass

        # normalize all text in this textbox
        full_text = "\n".join([p.text.strip() for p in paras if p.text.strip()])

        width = height = None
        parent = shape.getparent()
        if parent is not None:
            ext_elems = list(parent.iter(EXT))
            if ext_elems:
                cx = ext_elems[0].get("cx")
                cy = ext_elems[0].get("cy")
                width = emu_to_inches(cx)
                height = emu_to_inches(cy)

        # dedup key
        sig = (full_text, width, height)
        if sig in seen:
            continue
        seen.add(sig)

        textboxes.append((shape, paras, width, height))

    return textboxes


def compare_textboxes(doc1, doc2, output_doc):
    from difflib import SequenceMatcher

    output_doc.add_heading("Textbox Comparison", level=1)

    tb1 = extract_textboxes(doc1)
    tb2 = extract_textboxes(doc2)

    texts1 = ["\n".join(p.text.strip() for p in paras if p.text.strip()) for (_, paras, _, _) in tb1]
    texts2 = ["\n".join(p.text.strip() for p in paras if p.text.strip()) for (_, paras, _, _) in tb2]

    sm = SequenceMatcher(None, texts1, texts2)

    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            # check formatting-only or size-only diffs
            for k in range(0, min(i2 - i1, j2 - j1)):
                _, paras1, w1, h1 = tb1[i1 + k]
                _, paras2, w2, h2 = tb2[j1 + k]

                d1 = dominant_run_style(paras1[0]) if paras1 else {}
                d2 = dominant_run_style(paras2[0]) if paras2 else {}

                diffs = style_diffs(d1, d2)
                size_changed = (w1 != w2 or h1 != h2)

                if diffs or size_changed:
                    preview = texts2[j1 + k][:80]
                    output_doc.add_paragraph(f"[Modified Textbox {j1 + k + 1}] | Preview: \"{preview}\"", style="Heading 3")

                    for (key, a, b) in diffs:
                        label = f"[{key.replace('_',' ').title()} Changed]"
                        add_colored_paragraph(output_doc, label, f"{a} → {b}", (0, 0, 255))

                    if size_changed:
                        add_colored_paragraph(output_doc, "[Size Changed]", f"{w1}×{h1} → {w2}×{h2}", (0, 0, 255))

        elif tag == "replace":
            for ii, jj in zip(range(i1, i2), range(j1, j2)):
                old_text = texts1[ii]
                new_text = texts2[jj]

                # 🚨 check similarity — if too different, treat as remove + add
                sm_ratio = SequenceMatcher(None, old_text, new_text).ratio()
                if sm_ratio < 0.4:  # threshold (tweakable)
                    # Removed
                    output_doc.add_paragraph(f"[Removed Textbox {ii+1}] | Preview: \"{old_text[:80]}\"", style="Heading 3")
                    add_colored_paragraph(output_doc, "[Removed Textbox]", old_text, (255, 0, 0))

                    # Added
                    _, _, w, h = tb2[jj]
                    output_doc.add_paragraph(f"[Added Textbox {jj+1}] | Preview: \"{new_text[:80]}\"", style="Heading 3")
                    add_colored_paragraph(output_doc, "[Added Textbox]", new_text, (0, 128, 0))
                    if w and h:
                        add_colored_paragraph(output_doc, "[Size]", f"{w}×{h}", (0, 128, 0))
                    continue

                # Otherwise treat as modified
                _, paras1, w1, h1 = tb1[ii]
                _, paras2, w2, h2 = tb2[jj]

                preview = new_text[:80]
                output_doc.add_paragraph(f"[Modified Textbox {jj+1}] | Preview: \"{preview}\"", style="Heading 3")
                add_colored_paragraph(output_doc, "[Old Textbox]", old_text, (255, 0, 0))
                add_colored_paragraph(output_doc, "[New Textbox]", new_text, (0, 128, 0))

                for diff in get_word_diff(old_text, new_text) + detect_spacing_issues(old_text, new_text):
                    add_colored_paragraph(output_doc, "", diff, (255, 165, 0))

                d1 = dominant_run_style(paras1[0]) if paras1 else {}
                d2 = dominant_run_style(paras2[0]) if paras2 else {}
                for (key, a, b) in style_diffs(d1, d2):
                    add_colored_paragraph(output_doc, f"[{key.title()} Changed]", f"{a} → {b}", (0, 0, 255))

                if (w1 != w2 or h1 != h2):
                    add_colored_paragraph(output_doc, "[Size Changed]", f"{w1}×{h1} → {w2}×{h2}", (0, 0, 255))

        elif tag == "delete":
            for ii in range(i1, i2):
                old_text = texts1[ii]
                output_doc.add_paragraph(f"[Removed Textbox {ii+1}] | Preview: \"{old_text[:80]}\"", style="Heading 3")
                add_colored_paragraph(output_doc, "[Removed Textbox]", old_text, (255, 0, 0))

        elif tag == "insert":
            for jj in range(j1, j2):
                new_text = texts2[jj]
                _, _, w, h = tb2[jj]
                output_doc.add_paragraph(f"[Added Textbox {jj+1}] | Preview: \"{new_text[:80]}\"", style="Heading 3")
                add_colored_paragraph(output_doc, "[Added Textbox]", new_text, (0, 128, 0))
                if w and h:
                    add_colored_paragraph(output_doc, "[Size]", f"{w}×{h}", (0, 128, 0))



# ============================= Header and Footer =========================================
def compare_headers_footers(doc1, doc2, output_doc):
    output_doc.add_heading("Header/Footer Comparison", level=1)
    
    def get_section_paragraphs(doc, attr):
        items = []
        for idx, sec in enumerate(doc.sections):
            paras = getattr(sec, attr).paragraphs
            for p in paras:
                if p.text.strip():
                    items.append((idx + 1, p.text.strip()))
        return items

    headers1 = get_section_paragraphs(doc1, 'header')
    headers2 = get_section_paragraphs(doc2, 'header')
    footers1 = get_section_paragraphs(doc1, 'footer')
    footers2 = get_section_paragraphs(doc2, 'footer')

    def compare_parts(part1, part2, label):
        texts1 = [t[1] for t in part1]
        texts2 = [t[1] for t in part2]
        matcher = difflib.SequenceMatcher(None, texts1, texts2)
        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag == 'equal':
                continue
            elif tag == 'replace':
                for i, j in zip(range(i1, i2), range(j1, j2)):
                    sec1 = part1[i][0]
                    sec2 = part2[j][0]
                    preview = texts2[j][:80] + ("..." if len(texts2[j]) > 80 else "")
                    output_doc.add_paragraph(f"{label} - Section {sec2} | Preview: \"{preview}\"", style="Heading 3")
                    add_colored_paragraph(output_doc, f"[Old {label}]", texts1[i], (255, 0, 0))
                    add_colored_paragraph(output_doc, f"[New {label}]", texts2[j], (0, 128, 0))
                    for diff in get_word_diff(texts1[i], texts2[j]) + detect_spacing_issues(texts1[i], texts2[j]):
                        add_colored_paragraph(output_doc, "", diff, (255, 165, 0))
            elif tag == 'delete':
                removed = {}
                for i in range(i1, i2):
                    txt = texts1[i]
                    sec = part1[i][0]
                    removed.setdefault(txt, []).append(str(sec))
                for txt, secs in removed.items():
                    output_doc.add_paragraph(
    f"[Removed {label}] {txt}→ Removed from Sections: {', '.join(secs)}",
    style="Heading 3"
)
            elif tag == 'insert':
                for j in range(j1, j2):
                    sec = part2[j][0]
                    preview = texts2[j][:80] + ("..." if len(texts2[j]) > 80 else "")
                    output_doc.add_paragraph(f"Added {label} - Section {sec} | Preview: \"{preview}\"", style="Heading 3")
                    add_colored_paragraph(output_doc, f"[Added {label}]", texts2[j], (0, 128, 0))

    compare_parts(headers1, headers2, "Header")
    compare_parts(footers1, footers2, "Footer")

# ================================== Table Comparison =================================================

def get_cell_style(cell):
    style = {
        "text": cell.text.strip(),
        "font_name": "Default",
        "font_size": "Default",
        "bold": False,
        "italic": False,
        "underline": False,
        "highlight": None,
        "font_color": None,
        "alignment": "UNKNOWN",
        "spacing": "Default",
        "left_indent": 0.0,
        "right_indent": 0.0,
        "letter_spacing": "Default"
    }
    if cell.paragraphs:
        para = cell.paragraphs[0]
        run = para.runs[0] if para.runs else None
        pf = para.paragraph_format
        if run:
            spacing_val = get_letter_spacing_from_run(run)
            # Safe mappings
            highlight_val = run.font.highlight_color
            highlight_name = str(highlight_val).split('.')[-1] if highlight_val else None
            font_color_val = run.font.color.rgb
            font_color_name = None
            if font_color_val:
                font_color_name = str(font_color_val).lower()

            style.update({
                "font_name": run.font.name or "Default",
                "font_size": run.font.size.pt if run.font.size else "Default",
                "bold": True if run.bold else False,
                "italic": True if run.italic else False,
                "underline": True if run.underline else False,
                "highlight": highlight_name,
                "font_color": font_color_name,
                "alignment": get_alignment(para),
                "spacing": round(pf.line_spacing, 2) if pf.line_spacing else "Default",
                "left_indent": round(pf.left_indent.cm, 2) if pf.left_indent else 0.0,
                "right_indent": round(pf.right_indent.cm, 2) if pf.right_indent else 0.0,
                "letter_spacing": spacing_val
            })
    return style

def estimate_pages_with_breaks(doc):
    page_number = 1
    block_page_map = []
    for block in doc.element.body.iter():
        if block.tag.endswith("}br") and block.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type') == 'page':
            page_number += 1
        elif block.tag.endswith("}p"):
            block_page_map.append(("paragraph", page_number))
        elif block.tag.endswith("}tbl"):
            block_page_map.append(("table", page_number))
    return block_page_map

def get_table_estimated_page(doc, table_index):
    pages = estimate_pages_with_breaks(doc)
    table_count = -1
    for block_type, page in pages:
        if block_type == "table":
            table_count += 1
            if table_count == table_index:
                return page
    return 1

def compare_tables(doc1, doc2, output_doc):
    import difflib
    from docx.shared import RGBColor

    output_doc.add_heading("Table Comparison", level=1)

    tables1 = list(doc1.tables)
    tables2 = list(doc2.tables)

    # ---------- local helpers (self-contained) ----------
    def norm_space(s: str) -> str:
        return re.sub(r"\s+", " ", s or "").strip()

    def table_signature(tbl) -> str:
        parts = []
        for r in tbl.rows:
            row_parts = []
            for c in r.cells:
                row_parts.append(norm_space(c.text))
            parts.append("|".join(row_parts))
        return "||".join(parts)

    def table_dims(tbl):
        rows = len(tbl.rows)
        cols = len(tbl.rows[0].cells) if rows else 0
        return rows, cols

    def preview_text(tbl) -> str:
        try:
            for r in tbl.rows:
                for c in r.cells:
                    t = norm_space(c.text)
                    if t:
                        return t[:80]
        except Exception:
            pass
        return ""

    # color + highlight name helpers (no global deps)
    def rgb_name(rgb):
        if not rgb:
            return "None"
        try:
            r, g, b = rgb[0], rgb[1], rgb[2]
        except Exception:
            # python-docx RGBColor may be bytes-like or object; try attributes
            r, g, b = getattr(rgb, "r", None), getattr(rgb, "g", None), getattr(rgb, "b", None)
        if None in (r, g, b):
            return "None"
        known = {
            (255, 0, 0): "red",
            (0, 255, 0): "green",
            (0, 0, 255): "blue",
            (255, 255, 0): "yellow",
            (0, 0, 0): "black",
            (255, 255, 255): "white",
            (128, 128, 128): "gray",
            (0, 255, 255): "cyan",
            (255, 0, 255): "magenta",
        }
        name = known.get((r, g, b))
        return name if name else f"#{r:02X}{g:02X}{b:02X}"

    def highlight_name(h):
        from docx.enum.text import WD_COLOR_INDEX
        m = {
            WD_COLOR_INDEX.AUTO: "auto",
            WD_COLOR_INDEX.BLACK: "black",
            WD_COLOR_INDEX.BLUE: "blue",
            WD_COLOR_INDEX.BRIGHT_GREEN: "bright green",
            WD_COLOR_INDEX.DARK_BLUE: "dark blue",
            WD_COLOR_INDEX.DARK_RED: "dark red",
            WD_COLOR_INDEX.DARK_YELLOW: "dark yellow",
            WD_COLOR_INDEX.GRAY_25: "gray 25%",
            WD_COLOR_INDEX.GRAY_50: "gray 50%",
            WD_COLOR_INDEX.GREEN: "green",
            WD_COLOR_INDEX.PINK: "pink",
            WD_COLOR_INDEX.RED: "red",
            WD_COLOR_INDEX.TEAL: "teal",
            WD_COLOR_INDEX.TURQUOISE: "turquoise",
            WD_COLOR_INDEX.VIOLET: "violet",
            WD_COLOR_INDEX.WHITE: "white",
            WD_COLOR_INDEX.YELLOW: "yellow",
            None: "None",
        }
        return m.get(h, str(h))

    def to_bool(x):
        return True if x is True else False

    def first_para_and_run(cell):
        # pick the first paragraph that has at least one run; fall back to first paragraph
        if not cell.paragraphs:
            return None, None
        for p in cell.paragraphs:
            if p.runs:
                return p, p.runs[0]
        return cell.paragraphs[0], None

    def get_cell_style_safe(cell):
        style = {
            "font_name": "Default",
            "font_size": "Default",
            "bold": False,
            "italic": False,
            "underline": False,
            "font_color": "None",
            "highlight": "None",
            "alignment": "UNKNOWN",
            "spacing": "Default",
            "left_indent": 0.0,
            "right_indent": 0.0,
            "letter_spacing": "Default",
        }
        p, r = first_para_and_run(cell)
        if p:
            style["alignment"] = get_alignment(p)
            style["spacing"] = get_line_spacing(p)
            li, ri = get_indent(p)
            style["left_indent"] = li
            style["right_indent"] = ri
        if r:
            # font properties (normalize to avoid None noise)
            style["bold"] = to_bool(r.bold)
            style["italic"] = to_bool(r.italic)
            style["underline"] = to_bool(r.underline)
            style["font_name"] = r.font.name or "Default"
            style["font_size"] = r.font.size.pt if r.font.size else "Default"
            style["highlight"] = highlight_name(r.font.highlight_color if r.font else None)
            style["font_color"] = rgb_name(r.font.color.rgb if r.font and r.font.color and r.font.color.rgb else None)
            style["letter_spacing"] = get_letter_spacing_from_run(r)
        return style

    def similarity_score(sig1, sig2, dim1, dim2):
        # mix of text similarity + shape similarity
        text_sim = difflib.SequenceMatcher(None, sig1, sig2).ratio()
        rows1, cols1 = dim1
        rows2, cols2 = dim2
        shape_sim = 1.0 if (rows1 == rows2 and cols1 == cols2) else 0.6 if (rows1 == rows2 or cols1 == cols2) else 0.0
        return 0.85 * text_sim + 0.15 * shape_sim

    def build_match_map(tables1, tables2, threshold=0.65):
        # greedy best-match pairing with threshold
        sig1 = [table_signature(t) for t in tables1]
        sig2 = [table_signature(t) for t in tables2]
        dim1 = [table_dims(t) for t in tables1]
        dim2 = [table_dims(t) for t in tables2]

        pairs = []
        candidates = []
        for i in range(len(tables1)):
            for j in range(len(tables2)):
                sc = similarity_score(sig1[i], sig2[j], dim1[i], dim2[j])
                candidates.append((sc, i, j))
        # highest score first
        candidates.sort(reverse=True)

        used_old = set()
        used_new = set()
        for sc, i, j in candidates:
            if sc < threshold:
                break
            if i in used_old or j in used_new:
                continue
            used_old.add(i)
            used_new.add(j)
            pairs.append((i, j))

        old_to_new = {i: j for i, j in pairs}
        new_to_old = {j: i for i, j in pairs}
        unmatched_old = [i for i in range(len(tables1)) if i not in old_to_new]
        unmatched_new = [j for j in range(len(tables2)) if j not in new_to_old]
        return old_to_new, new_to_old, unmatched_old, unmatched_new
    # ---------- end helpers ----------

    old_to_new, new_to_old, unmatched_old, unmatched_new = build_match_map(tables1, tables2)

    i = 0  # index in old
    j = 0  # index in new
    n_old = len(tables1)
    n_new = len(tables2)

    while i < n_old or j < n_new:
        # handle removed tables (in order)
        if i < n_old and i not in old_to_new:
            t1 = tables1[i]
            page_est = get_table_estimated_page(doc1, i)
            prev1 = preview_text(t1)
            p = output_doc.add_paragraph(f"[Removed Table {i+1}] | Page {page_est} | Preview: \"{prev1}\"")
            p.runs[0].font.color.rgb = RGBColor(255, 0, 0)
            i += 1
            continue

        # handle added tables that appear before the next matched pair
        if j < n_new and j not in new_to_old:
            t2 = tables2[j]
            page_est = get_table_estimated_page(doc2, j)
            prev2 = preview_text(t2)
            p = output_doc.add_paragraph(f"[Added Table {j+1}] | Page {page_est} | Preview: \"{prev2}\"")
            p.runs[0].font.color.rgb = RGBColor(0, 128, 0)
            j += 1
            continue

        # matched pair at positions i and j
        if i < n_old and j < n_new and old_to_new.get(i) == j:
            t1 = tables1[i]
            t2 = tables2[j]
            page_est = get_table_estimated_page(doc2, j)
            prev2 = preview_text(t2)

            table_changes = []

            # shape changes (rows/cols added/removed)
            r1, c1 = table_dims(t1)
            r2, c2 = table_dims(t2)
            if r1 != r2:
                table_changes.append(f"[Rows Changed] {r1} → {r2}")
            if c1 != c2:
                table_changes.append(f"[Columns Changed] {c1} → {c2}")

            # compare only intersection to avoid index errors
            min_rows = min(r1, r2)
            min_cols = min(c1, c2)

            for r_idx in range(min_rows):
                for c_idx in range(min_cols):
                    # ✅ SAFETY GUARD for cell access
                    try:
                        cell1 = t1.cell(r_idx, c_idx)
                        cell2 = t2.cell(r_idx, c_idx)
                    except IndexError:
                        continue

                    text1 = norm_space(cell1.text)
                    text2 = norm_space(cell2.text)

                    if text1 != text2:
                        # textual diffs + spacing diagnostics
                        table_changes.append(f"[Cell ({r_idx+1},{c_idx+1}) Old] {text1}")
                        table_changes.append(f"[Cell ({r_idx+1},{c_idx+1}) New] {text2}")
                        for diff in get_word_diff(text1, text2) + detect_spacing_issues(text1, text2):
                            table_changes.append(f"  {diff}")
                    else:
                        # only if text equal, check formatting deltas (safe, minimal noise)
                        s1 = get_cell_style_safe(cell1)
                        s2 = get_cell_style_safe(cell2)
                        for key in [
                            "alignment", "spacing", "left_indent", "right_indent",
                            "font_name", "font_size", "bold", "italic", "underline",
                            "font_color", "highlight", "letter_spacing"
                        ]:
                            if s1.get(key) != s2.get(key):
                                table_changes.append(f"  [{key.replace('_',' ').title()} Changed] {s1.get(key)} → {s2.get(key)}")

            if table_changes:
                output_doc.add_paragraph(f"[Modified Table {j+1}] | Page {page_est} | Preview: \"{prev2}\"", style="Heading 3")
                for ch in table_changes:
                    if ch.startswith("[Cell") and " Old]" in ch:
                        add_colored_paragraph(output_doc, "", ch, (255, 0, 0))
                    elif ch.startswith("[Cell") and " New]" in ch:
                        add_colored_paragraph(output_doc, "", ch, (0, 128, 0))
                    else:
                        add_colored_paragraph(output_doc, "", ch, (255, 165, 0))

            i += 1
            j += 1
            continue

        # If we get here, indices are out of sync; advance safely
        if i < n_old and j < n_new:
            # move forward to the next mapped slot
            target_j = old_to_new.get(i, None)
            if target_j is not None and target_j > j:
                # there are added tables before the mapped one
                t2 = tables2[j]
                page_est = get_table_estimated_page(doc2, j)
                prev2 = preview_text(t2)
                p = output_doc.add_paragraph(f"[Added Table {j+1}] | Page {page_est} | Preview: \"{prev2}\"")
                p.runs[0].font.color.rgb = RGBColor(0, 128, 0)
                j += 1
            else:
                # treat the old table as removed
                t1 = tables1[i]
                page_est = get_table_estimated_page(doc1, i)
                prev1 = preview_text(t1)
                p = output_doc.add_paragraph(f"[Removed Table {i+1}] | Page {page_est} | Preview: \"{prev1}\"")
                p.runs[0].font.color.rgb = RGBColor(255, 0, 0)
                i += 1

# =============================== Image Comparison ========================================
EMU_PER_INCH = 914400

def emu_to_inches_img(emu):
    try:
        return round(int(emu) / EMU_PER_INCH, 2)
    except Exception:
        return None


def extract_images(doc):
    
    images = []
    for i, shape in enumerate(doc.inline_shapes, start=1):
        try:
            width = emu_to_inches_img(shape._inline.extent.cx)
            height = emu_to_inches_img(shape._inline.extent.cy)

            # Find parent paragraph safely
            p = shape._inline.getparent()
            while p is not None and p.tag != qn("w:p"):
                p = p.getparent()

            para_index = None
            if p is not None:
                for idx, para in enumerate(doc.paragraphs, start=1):
                    if para._p == p:
                        para_index = idx
                        break

            images.append({
                "id": i,
                "width": width,
                "height": height,
                "para_index": para_index
            })
        except Exception:
            continue
    return images

def compare_images(doc1, doc2, output_doc):
    
    import hashlib, io, math
    from collections import defaultdict

    EMU_PER_INCH_IMG = 914400
    NS = {
        'w':  "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        'wp': "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        'a':  "http://schemas.openxmlformats.org/drawingml/2006/main",
        'pic':"http://schemas.openxmlformats.org/drawingml/2006/picture",
        'r':  "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    }

    # thresholds (tweakable)
    VHASH_HAMMING_THRESHOLD = 10      # for 64-bit aHash (<=10 considered same-ish)
    SIZE_MATCH_TOLERANCE_IN = 0.6     # inches total tolerance (w diff + h diff)
    MIN_SIZE_FILTER_IN = 0.0          # ignore images smaller than this (both dims); 0 = keep all

    def emu_to_inches_img(emu_val):
        try:
            return round(int(emu_val) / EMU_PER_INCH_IMG, 2)
        except Exception:
            return None

    # compute hamming distance
    def hamming(a, b):
        if a is None or b is None:
            return 999
        x = a ^ b
        return x.bit_count() if hasattr(x, "bit_count") else bin(x).count("1")

    # Robust extractor: walks paragraph runs and finds a:blip r:embed
    def extract_images_for_compare(doc):
        imgs = []
        idx = 0
        try:
            from PIL import Image
            pil_ok = True
        except Exception:
            pil_ok = False
            Image = None

        for p_idx, para in enumerate(doc.paragraphs, start=1):
            for run in para.runs:
                r = getattr(run, "_r", None)
                if r is None:
                    continue

                # try xpath for blips (works in most python-docx versions)
                try:
                    blip_nodes = r.xpath('.//a:blip', namespaces=NS)
                except Exception:
                    # fallback: iterate and find elements ending with '}blip'
                    blip_nodes = [el for el in r.iter() if str(el.tag).endswith('}blip')]

                if not blip_nodes:
                    continue

                for blip in blip_nodes:
                    rid = blip.get('{%s}embed' % NS['r'])
                    if not rid:
                        continue
                    # get part blob
                    try:
                        part = doc.part.related_parts[rid]
                        blob = part.blob
                    except Exception:
                        # relationship missing; skip
                        continue

                    # get size if present in drawing extent (wp:extent)
                    w_in = h_in = None
                    try:
                        ext_nodes = r.xpath('.//wp:extent', namespaces=NS)
                    except Exception:
                        ext_nodes = [el for el in r.iter() if str(el.tag).endswith('}extent')]
                    if ext_nodes:
                        cx = ext_nodes[0].get('cx')
                        cy = ext_nodes[0].get('cy')
                        w_in = emu_to_inches_img(cx) if cx else None
                        h_in = emu_to_inches_img(cy) if cy else None

                    # optional tiny filtering
                    if (w_in is not None and h_in is not None and
                        (w_in < MIN_SIZE_FILTER_IN and h_in < MIN_SIZE_FILTER_IN)):
                        continue

                    # content hash
                    sha1 = hashlib.sha1(blob).hexdigest()

                    # visual hash (aHash 8x8) if Pillow available
                    vhash = None
                    if pil_ok:
                        try:
                            img = Image.open(io.BytesIO(blob)).convert('L')
                            # resize to 8x8
                            if hasattr(Image, "Resampling"):
                                img = img.resize((8, 8), Image.Resampling.LANCZOS)
                            else:
                                img = img.resize((8, 8), Image.ANTIALIAS)
                            pixels = list(img.getdata())
                            avg = sum(pixels) // len(pixels)
                            bits = 0
                            for px in pixels:
                                bits = (bits << 1) | (1 if px > avg else 0)
                            vhash = bits  # 64-bit-ish integer
                        except Exception:
                            vhash = None

                    idx += 1
                    imgs.append({
                        "sha1": sha1,
                        "vhash": vhash,
                        "w": w_in,
                        "h": h_in,
                        "pidx": p_idx,
                        "order_idx": idx,
                        "blob": blob
                    })
        return imgs

    # pairing helpers
    def greedy_match_by_key(list_a, list_b, key_fn, score_fn, threshold):
       
        A = list(range(len(list_a)))
        B = list(range(len(list_b)))
        candidates = []
        for i in A:
            for j in B:
                s = score_fn(list_a[i], list_b[j])
                if s >= threshold:
                    candidates.append((s, i, j))
        # sort descending by score
        candidates.sort(reverse=True, key=lambda x: x[0])
        usedA, usedB = set(), set()
        pairs = []
        for s, i, j in candidates:
            if i in usedA or j in usedB:
                continue
            usedA.add(i); usedB.add(j)
            pairs.append((i, j))
        leftoverA = [i for i in A if i not in usedA]
        leftoverB = [j for j in B if j not in usedB]
        return pairs, leftoverA, leftoverB

    # score functions
    def score_sha(a, b):
        return 1.0 if a["sha1"] == b["sha1"] else 0.0

    def score_vhash(a, b):
        if a["vhash"] is None or b["vhash"] is None:
            return 0.0
        # score = inverse normalized hamming distance (64 bits)
        ham = hamming(a["vhash"], b["vhash"])
        # map ham 0..64 to score 1..0
        sc = max(0.0, 1.0 - (ham / 64.0))
        return sc

    def score_size_and_proximity(a, b):
        # size similarity: smaller diff -> higher score
        aw, ah = a.get("w") or 0, a.get("h") or 0
        bw, bh = b.get("w") or 0, b.get("h") or 0
        size_diff = abs(aw - bw) + abs(ah - bh)
        size_score = max(0.0, 1.0 - (size_diff / SIZE_MATCH_TOLERANCE_IN))
        # paragraph proximity
        pa, pb = a.get("pidx") or 0, b.get("pidx") or 0
        prox = abs((pa - pb))
        prox_score = max(0.0, 1.0 - min(prox, 20) / 20.0)
        return 0.6 * size_score + 0.4 * prox_score

    # run extraction
    imgs1 = extract_images_for_compare(doc1)
    imgs2 = extract_images_for_compare(doc2)

    output_doc.add_heading("Image Comparison", level=1)

    # 1) exact sha matches (score 1.0)
    pairs_sha, rem1, rem2 = greedy_match_by_key(imgs1, imgs2, lambda x: x["sha1"], lambda a,b: 1.0 if a["sha1"] == b["sha1"] else 0.0, 1.0)

    matched_a = set(i for i,_ in pairs_sha)
    matched_b = set(j for _,j in pairs_sha)
    final_pairs = list(pairs_sha)

    # 2) visual hash matching for leftovers if Pillow available
    leftovers1 = [imgs1[i] for i in range(len(imgs1)) if i not in matched_a]
    leftovers2 = [imgs2[j] for j in range(len(imgs2)) if j not in matched_b]
    if leftovers1 and leftovers2:
        pairs_v, la, lb = greedy_match_by_key(leftovers1, leftovers2, None, score_vhash, 1.0 - (VHASH_HAMMING_THRESHOLD / 64.0))
        # map back to original indices
        # leftovers1[i] corresponds to original index idx1 where imgs1[idx1] == leftovers1[i]
        idx_map1 = [i for i in range(len(imgs1)) if i not in matched_a]
        idx_map2 = [j for j in range(len(imgs2)) if j not in matched_b]
        for i_local, j_local in pairs_v:
            i_orig = idx_map1[i_local]
            j_orig = idx_map2[j_local]
            final_pairs.append((i_orig, j_orig))
            matched_a.add(i_orig); matched_b.add(j_orig)

    # 3) size+proximity match for remaining
    leftovers1 = [i for i in range(len(imgs1)) if i not in matched_a]
    leftovers2 = [j for j in range(len(imgs2)) if j not in matched_b]
    if leftovers1 and leftovers2:
        listA = [imgs1[i] for i in leftovers1]
        listB = [imgs2[j] for j in leftovers2]
        # use score_size_and_proximity, threshold low
        pairs_s, la, lb = greedy_match_by_key(listA, listB, None, score_size_and_proximity, 0.25)
        # map back indices
        for i_local, j_local in pairs_s:
            i_orig = leftovers1[i_local]
            j_orig = leftovers2[j_local]
            final_pairs.append((i_orig, j_orig))
            matched_a.add(i_orig); matched_b.add(j_orig)

    # Now produce reports
    used_pairs = set()
    counter = 1
    # sort pairs by new doc order for nicer output
    final_pairs_sorted = sorted(final_pairs, key=lambda pair: imgs2[pair[1]]["order_idx"] if pair[1] < len(imgs2) else 0)
    for (ia, jb) in final_pairs_sorted:
        a = imgs1[ia]; b = imgs2[jb]
        moved = (a.get("pidx") != b.get("pidx"))
        size_changed = (a.get("w") != b.get("w") or a.get("h") != b.get("h"))

        if size_changed and not moved:
            add_colored_paragraph(
                output_doc,
                f"[Image Modified] Image {counter}",
                f"Old: Size={a.get('w')}in × {a.get('h')}in\nNew: Size={b.get('w')}in × {b.get('h')}in",
                (0, 0, 255)
            )
        elif moved and not size_changed:
            add_colored_paragraph(
                output_doc,
                f"[Image Moved] Image {counter}",
                f"Size: {b.get('w')}in × {b.get('h')}in\nOld position: Paragraph {a.get('pidx')} → New position: Paragraph {b.get('pidx')}",
                (255, 165, 0)
            )
        elif moved and size_changed:
            add_colored_paragraph(
                output_doc,
                f"[Image Modified] Image {counter}",
                f"Old: Size={a.get('w')}in × {a.get('h')}in\nNew: Size={b.get('w')}in × {b.get('h')}in\nOld position: Paragraph {a.get('pidx')} → New: Paragraph {b.get('pidx')}",
                (0, 0, 255)
            )
        else:
            # exact same content and size/position -> nothing to report
            pass

        used_pairs.add((ia, jb))
        counter += 1

    # Removed (present in doc1 but not matched)
    for i, rec in enumerate(imgs1):
        if i in matched_a:
            continue
        add_colored_paragraph(
            output_doc,
            f"[Image Removed] Image {counter}",
            f"Size: {rec.get('w')}in × {rec.get('h')}in\nOld position: Paragraph {rec.get('pidx')}",
            (255, 0, 0)
        )
        counter += 1

    # Added (present in doc2 but not matched)
    for j, rec in enumerate(imgs2):
        if j in matched_b:
            continue
        add_colored_paragraph(
            output_doc,
            f"[Image Added] Image {counter}",
            f"Size: {rec.get('w')}in × {rec.get('h')}in\nPosition: Paragraph {rec.get('pidx')}",
            (0, 128, 0)
        )
        counter += 1


# =================================Shape Comparison ========================================================
EMU_PER_INCH = 914400

def emu_to_inches(emu):
    return round(emu / EMU_PER_INCH, 2)

def extract_shapes_from_doc(doc):
    shapes = []
    package = doc._part.package
    document_part = package.part_related_by(
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument"
    )
    xml_content = document_part.blob
    tree = ET.fromstring(xml_content)

    ns = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
        'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape'
    }

    for drawing in tree.findall(".//w:drawing", ns):
        # 🚫 Skip if it's an image
        if drawing.find(".//pic:pic", ns) is not None:
            continue

        extent = drawing.find(".//wp:extent", ns)
        if extent is None:
            continue
        cx = int(extent.attrib.get('cx', 0))
        cy = int(extent.attrib.get('cy', 0))

        # ✅ Distinguish textbox vs real shape
        if drawing.find(".//wps:txbx", ns) is not None:
            shape_type = "textbox"
        else:
            prstGeom = drawing.find(".//a:prstGeom", ns)
            shape_type = prstGeom.attrib.get('prst') if prstGeom is not None else "unknown"

        shapes.append({
            "type": shape_type,
            "width": round(cx / EMU_PER_INCH, 2),
            "height": round(cy / EMU_PER_INCH, 2)
        })

    return shapes


def shape_distance(s1, s2):
    import math
    type_score = 0 if s1['type'] == s2['type'] else 1
    size_score = math.sqrt((s1['width'] - s2['width'])**2 + (s1['height'] - s2['height'])**2)
    return type_score * 5 + size_score

def compare_shapes(doc1, doc2, output_doc):
   
    import math
    from docx.shared import RGBColor

    def size_dist(a, b):
        return math.hypot(a['width'] - b['width'], a['height'] - b['height'])

    # --- pull shapes you already extract ---
    shapes1 = extract_shapes_from_doc(doc1)  # expects dicts: {'type','width','height'}
    shapes2 = extract_shapes_from_doc(doc2)

    output_doc.add_heading("Shape Comparison", level=1)

    n, m = len(shapes1), len(shapes2)
    if n == 0 and m == 0:
        return

    # --- costs ---
    ADD_COST = 3.0
    REM_COST = 3.0
    BIG = 1e6  # disallow cross-type matches

    def match_cost(s1, s2):
        if s1['type'] != s2['type']:
            return BIG  # never force different types to match
        return size_dist(s1, s2)

    # --- DP (Needleman–Wunsch style) ---
    dp = [[0.0] * (m + 1) for _ in range(n + 1)]
    bt = [[None] * (m + 1) for _ in range(n + 1)]  # backtrace

    # init
    for i in range(1, n + 1):
        dp[i][0] = dp[i - 1][0] + REM_COST
        bt[i][0] = ('REM', i - 1, None)
    for j in range(1, m + 1):
        dp[0][j] = dp[0][j - 1] + ADD_COST
        bt[0][j] = ('ADD', None, j - 1)

    # fill
    for i in range(1, n + 1):
        for j in range(1, m + 1):
            c_match = dp[i - 1][j - 1] + match_cost(shapes1[i - 1], shapes2[j - 1])
            c_rem   = dp[i - 1][j] + REM_COST
            c_add   = dp[i][j - 1] + ADD_COST

            # pick min
            if c_match <= c_rem and c_match <= c_add:
                dp[i][j] = c_match
                bt[i][j] = ('MATCH', i - 1, j - 1)
            elif c_rem <= c_add:
                dp[i][j] = c_rem
                bt[i][j] = ('REM', i - 1, None)
            else:
                dp[i][j] = c_add
                bt[i][j] = ('ADD', None, j - 1)

    # backtrack to actions in forward order
    actions = []
    i, j = n, m
    while i > 0 or j > 0:
        op, oi, oj = bt[i][j]
        actions.append((op, oi, oj))
        if op == 'MATCH':
            i -= 1; j -= 1
        elif op == 'REM':
            i -= 1
        else:  # 'ADD'
            j -= 1
    actions.reverse()

    # --- emit results (suppress tiny noise) ---
    EPS = 0.05  # ~0.05 inch = ignore tiny differences
    shape_counter = 1

    for op, oi, oj in actions:
        if op == 'MATCH':
            s1 = shapes1[oi]
            s2 = shapes2[oj]
            # same type guaranteed by match_cost
            if size_dist(s1, s2) > EPS:
                p = output_doc.add_paragraph(f"[Shape Modified] Shape {shape_counter}")
                p.runs[0].font.color.rgb = RGBColor(0, 0, 255)
                output_doc.add_paragraph(f"Old: Type={s1['type']}, W={s1['width']}in, H={s1['height']}in")
                output_doc.add_paragraph(f"New: Type={s2['type']}, W={s2['width']}in, H={s2['height']}in")
            # else identical enough → no line
            shape_counter += 1

        elif op == 'REM':
            s1 = shapes1[oi]
            p = output_doc.add_paragraph(f"[Shape Removed] Shape {shape_counter}")
            p.runs[0].font.color.rgb = RGBColor(255, 0, 0)
            output_doc.add_paragraph(f"Old: Type={s1['type']}, W={s1['width']}in, H={s1['height']}in")
            output_doc.add_paragraph("New: None")
            shape_counter += 1

        elif op == 'ADD':
            s2 = shapes2[oj]
            p = output_doc.add_paragraph(f"[Shape Added] Shape {shape_counter}")
            p.runs[0].font.color.rgb = RGBColor(0, 128, 0)
            output_doc.add_paragraph("Old: None")
            output_doc.add_paragraph(f"New: Type={s2['type']}, W={s2['width']}in, H={s2['height']}in")
            shape_counter += 1
        
#================================ Page Break Comparison ========================================

def extract_page_break_positions(doc):
    
    page_breaks = []
    for i, para in enumerate(doc.paragraphs):
        for run in para.runs:
            # Look for <w:br w:type="page"/> in the run XML
            for br in run._element.findall(".//w:br", run._element.nsmap):
                if br.get(qn("w:type")) == "page":
                    preview = para.text.strip()[:80]
                    page_breaks.append((i + 1, preview))
    return page_breaks


def compare_page_breaks(doc1, doc2, output_doc):
    output_doc.add_heading("Page Break Comparison", level=1)

    pb1 = extract_page_break_positions(doc1)
    pb2 = extract_page_break_positions(doc2)

    i, j = 0, 0
    changes_found = False

    while i < len(pb1) and j < len(pb2):
        if pb1[i] == pb2[j]:
            i += 1
            j += 1
        else:
            p1, prev1 = pb1[i]
            p2, prev2 = pb2[j]
            if p1 < p2:
                p = output_doc.add_paragraph(
                    f"[Removed PAGE_BREAK] after Paragraph {p1} | Preview: \"{prev1}\""
                )
                p.runs[0].font.color.rgb = RGBColor(255, 0, 0)
                i += 1
                changes_found = True
            elif p2 < p1:
                p = output_doc.add_paragraph(
                    f"[Added PAGE_BREAK] after Paragraph {p2} | Preview: \"{prev2}\""
                )
                p.runs[0].font.color.rgb = RGBColor(0, 128, 0)
                j += 1
                changes_found = True
            else:
                p = output_doc.add_paragraph(
                    f"[Modified PAGE_BREAK] Paragraph {p1}"
                )
                p.runs[0].font.color.rgb = RGBColor(0, 0, 255)
                i += 1
                j += 1
                changes_found = True

    # Remaining removals
    while i < len(pb1):
        p1, prev1 = pb1[i]
        p = output_doc.add_paragraph(
            f"[Removed PAGE_BREAK] after Paragraph {p1} | Preview: \"{prev1}\""
        )
        p.runs[0].font.color.rgb = RGBColor(255, 0, 0)
        i += 1
        changes_found = True

    # Remaining additions
    while j < len(pb2):
        p2, prev2 = pb2[j]
        p = output_doc.add_paragraph(
            f"[Added PAGE_BREAK] after Paragraph {p2} | Preview: \"{prev2}\""
        )
        p.runs[0].font.color.rgb = RGBColor(0, 128, 0)
        j += 1
        changes_found = True

    if not changes_found:
        output_doc.add_paragraph("No page break differences found ✅")

#===================================== Moved paragraph comparison ==================================

from collections import defaultdict

def group_paragraphs(doc, block_size=1):
    """Group paragraphs into blocks (default 1 para per block)."""
    blocks = []
    buf = []
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt:
            buf.append(txt)
            if len(buf) >= block_size:
                blocks.append("\n".join(buf))
                buf = []
    if buf:
        blocks.append("\n".join(buf))
    return blocks

def compare_moved_paragraphs(doc1, doc2, output_doc):
    output_doc.add_heading("Moved Paragraphs Comparison", level=1)

    blocks1 = group_paragraphs(doc1, block_size=1)
    blocks2 = group_paragraphs(doc2, block_size=1)

    pos1 = defaultdict(list)
    pos2 = defaultdict(list)

    for i, b in enumerate(blocks1):
        pos1[b].append(i)
    for j, b in enumerate(blocks2):
        pos2[b].append(j)

    found = False
    reported = set()  # to avoid duplicates

    for text in set(pos1.keys()).intersection(pos2.keys()):
        # only check the first occurrence for simplicity
        idx1 = pos1[text][0]
        idx2 = pos2[text][0]

        if idx1 != idx2:
            key = (text, min(idx1, idx2), max(idx1, idx2))
            if key in reported:
                continue
            reported.add(key)

            found = True
            preview = text[:120].replace("\n", " / ")
            output_doc.add_paragraph(
                f"[Paragraph Block Moved] \"{preview}\"",
                style="Heading 3"
            )
            output_doc.add_paragraph(
                f"Old position: Block {idx1+1} → New position: Block {idx2+1}"
            )

    if not found:
        output_doc.add_paragraph("No moved paragraphs found ✅")

# ====================== SREAMLIT UI ================================
if __name__ == "__main__":
    st.set_page_config(page_title="Word File Comparator", layout="centered")
    st.title("📄 Word Document Comparator")
    st.markdown("Upload your **Pre** and **Post** Word documents to compare changes like **text, formatting, spacing, fonts, images, tables, headers/footers, etc.**")

    st.subheader("📤 Upload Pre Document (.docx)")
    uploaded_pre = st.file_uploader("Upload the PRE document", type=["docx"], key="pre")

    st.subheader("📥 Upload Post Document (.docx)")
    uploaded_post = st.file_uploader("Upload the POST document", type=["docx"], key="post")

    if uploaded_pre and uploaded_post:
        if st.button("🔍 Compare Documents"):
            with st.spinner("Comparing documents, please wait..."):
                doc1 = Document(uploaded_pre)
                doc2 = Document(uploaded_post)
                output_doc = Document()

                compare_moved_paragraphs(doc1, doc2, output_doc)
                compare_paragraphs(doc1, doc2, output_doc)
                compare_textboxes(doc1, doc2, output_doc)
                compare_headers_footers(doc1, doc2, output_doc)
                compare_tables(doc1, doc2, output_doc)
                compare_images(doc1, doc2, output_doc)
                compare_shapes(doc1, doc2, output_doc)
                compare_page_breaks(doc1, doc2, output_doc)
                

                buffer = BytesIO()
                output_doc.save(buffer)
                buffer.seek(0)

                pre_name = uploaded_pre.name.rsplit('.', 1)[0]
                post_name = uploaded_post.name.rsplit('.', 1)[0]
                dynamic_filename = f"Comparison_{pre_name}_vs_{post_name}.docx"

            st.success("✅ Comparison Complete!")

            st.markdown("### 📊 Summary of Changes")
            st.write("- Moved Paragraphs compared ✅")
            st.write("- Paragraphs compared ✅")
            st.write("- Tables compared ✅")
            st.write("- Images compared ✅")
            st.write("- Text boxes compared ✅")
            st.write("- Headers/Footers compared ✅")
            st.write("- Shape Compared ✅")
            st.write("- Page Break Compared ✅")

            st.download_button(
                "⬇️ Download Comparison Report",
                data=buffer.getvalue(),
                file_name=dynamic_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    else:
        st.info("📌 Please upload both **Pre** and **Post** documents to begin comparison.")


